from docx import Document

def build_docx(path, data):

    doc = Document()

    # Header
    if data.get("header_text"):
        doc.add_heading(data["header_text"], level=1)

    doc.add_heading(data.get("chapter", "Lesson"), level=2)

    # Summary
    doc.add_heading("Summary", level=3)

    summary = data.get("summary_section", {})
    for line in summary.get("summary_lines", []):
        doc.add_paragraph(line)

    # True / False
    doc.add_heading("True or False", level=3)

    for q, a in data.get("true_false", []):
        doc.add_paragraph(q)

    # Fill blanks
    doc.add_heading("Fill in the Blanks", level=3)

    for q, _ in data.get("fill_blanks", []):
        doc.add_paragraph(q)

    # MCQ
    doc.add_heading("Multiple Choice Questions", level=3)

    for q, options, _ in data.get("mcqs", []):
        doc.add_paragraph(q)

        for opt in options:
            doc.add_paragraph(opt, style="List Bullet")

    doc.save(path)